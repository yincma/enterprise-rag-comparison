"""
文档处理模块
负责处理各种格式的文档，包括PDF、Word、Markdown、TXT等
"""

import os
import logging
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
import hashlib

# 文档处理库
import PyPDF2
import pdfplumber
from docx import Document
import markdown
from bs4 import BeautifulSoup

# 本地模块
import sys
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from utils.config import config_manager
from utils.helpers import (
    get_file_hash, get_file_size, format_file_size, 
    clean_text, split_text_into_chunks, validate_file_type,
    measure_performance, Timer
)
from utils.memory_optimizer import memory_optimized, batch_processor, memory_optimizer

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """文档处理器"""
    
    def __init__(self):
        """初始化文档处理器"""
        self.config = config_manager.load_app_config()
        self.doc_config = self.config.get('document_processing', {})
        
        # 支持的文件格式
        self.supported_formats = self.doc_config.get('supported_formats', ['.pdf', '.docx', '.txt', '.md'])
        
        # 文件大小限制
        self.max_file_size = self.doc_config.get('max_file_size', 104857600)  # 100MB
        
        # 分块配置
        self.chunk_config = self.config.get('vector_store', {})
        self.chunk_size = self.chunk_config.get('chunk_size', 1000)
        self.chunk_overlap = self.chunk_config.get('chunk_overlap', 200)
        
        logger.info("文档处理器初始化完成")
    
    @measure_performance
    def process_document(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        处理单个文档
        
        Args:
            file_path: 文档文件路径
            
        Returns:
            处理结果字典，包含文档内容和元数据
        """
        file_path = Path(file_path)
        
        # 验证文件
        if not self._validate_file(file_path):
            raise ValueError(f"文件验证失败: {file_path}")
        
        with Timer(f"处理文档 {file_path.name}"):
            # 提取文本内容
            text_content = self._extract_text(file_path)
            
            # 清理文本
            cleaned_text = clean_text(text_content)
            
            # 分割文本块
            text_chunks = split_text_into_chunks(
                cleaned_text,
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap
            )
            
            # 生成文档元数据
            metadata = self._generate_metadata(file_path, len(text_chunks))
            
            result = {
                'file_path': str(file_path),
                'filename': file_path.name,
                'text_content': cleaned_text,
                'text_chunks': text_chunks,
                'metadata': metadata,
                'chunk_count': len(text_chunks),
                'total_length': len(cleaned_text)
            }
            
            logger.info(f"文档处理完成: {file_path.name}, 生成 {len(text_chunks)} 个文本块")
            return result
    
    @batch_processor(batch_size=5, memory_limit_mb=1024)  # 每批处理5个文档，内存限制1GB
    def process_multiple_documents(self, file_paths: List[Union[str, Path]]) -> List[Dict[str, Any]]:
        """
        批量处理多个文档
        
        Args:
            file_paths: 文档文件路径列表
            
        Returns:
            处理结果列表
        """
        results = []
        successful_count = 0
        failed_count = 0
        
        logger.info(f"开始批量处理 {len(file_paths)} 个文档")
        
        for file_path in file_paths:
            try:
                result = self.process_document(file_path)
                results.append(result)
                successful_count += 1
            except Exception as e:
                logger.error(f"处理文档失败: {file_path}, 错误: {e}")
                failed_count += 1
                # 添加失败记录
                results.append({
                    'file_path': str(file_path),
                    'filename': Path(file_path).name,
                    'error': str(e),
                    'status': 'failed'
                })
        
        logger.info(f"批量处理完成: 成功 {successful_count} 个, 失败 {failed_count} 个")
        return results
    
    @memory_optimized(cache_name="document_text_cache", max_size=100, ttl=3600)  # 缓存1小时
    def _extract_text(self, file_path: Path) -> str:
        """
        根据文件类型提取文本内容
        
        Args:
            file_path: 文件路径
            
        Returns:
            提取的文本内容
        """
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.pdf':
            return self._extract_pdf_text(file_path)
        elif file_extension == '.docx':
            return self._extract_docx_text(file_path)
        elif file_extension == '.txt':
            return self._extract_txt_text(file_path)
        elif file_extension == '.md':
            return self._extract_markdown_text(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {file_extension}")
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """提取PDF文本内容"""
        text = ""
        
        try:
            # 首先尝试使用pdfplumber（更好的文本提取）
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if not text.strip():
                # 如果pdfplumber失败，尝试PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        text += page.extract_text() + "\n"
        
        except Exception as e:
            logger.error(f"PDF文本提取失败: {file_path}, 错误: {e}")
            raise
        
        return text
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """提取Word文档文本内容"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # 如果配置允许，提取表格文本
            if self.doc_config.get('docx', {}).get('extract_tables', True):
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts)
        
        except Exception as e:
            logger.error(f"Word文档文本提取失败: {file_path}, 错误: {e}")
            raise
    
    def _extract_txt_text(self, file_path: Path) -> str:
        """提取纯文本文件内容"""
        try:
            # 尝试不同的编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            # 如果所有编码都失败，使用错误忽略模式
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read()
        
        except Exception as e:
            logger.error(f"文本文件读取失败: {file_path}, 错误: {e}")
            raise
    
    def _extract_markdown_text(self, file_path: Path) -> str:
        """提取Markdown文件文本内容"""
        try:
            # 读取Markdown文件
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            # 将Markdown转换为HTML
            html = markdown.markdown(md_content)
            
            # 从HTML中提取纯文本
            soup = BeautifulSoup(html, 'html.parser')
            text = soup.get_text()
            
            return text
        
        except Exception as e:
            logger.error(f"Markdown文件处理失败: {file_path}, 错误: {e}")
            raise
    
    def _validate_file(self, file_path: Path) -> bool:
        """
        验证文件是否有效
        
        Args:
            file_path: 文件路径
            
        Returns:
            是否有效
        """
        try:
            # 检查文件是否存在
            if not file_path.exists():
                logger.error(f"文件不存在: {file_path}")
                return False
            
            # 检查是否是文件
            if not file_path.is_file():
                logger.error(f"路径不是文件: {file_path}")
                return False
            
            # 检查文件格式
            if not validate_file_type(file_path, self.supported_formats):
                logger.error(f"不支持的文件格式: {file_path}")
                return False
            
            # 检查文件大小
            file_size = get_file_size(file_path)
            if file_size > self.max_file_size:
                logger.error(f"文件过大: {file_path}, 大小: {format_file_size(file_size)}")
                return False
            
            return True
        
        except Exception as e:
            logger.error(f"文件验证失败: {file_path}, 错误: {e}")
            return False
    
    def _generate_metadata(self, file_path: Path, chunk_count: int) -> Dict[str, Any]:
        """
        生成文档元数据
        
        Args:
            file_path: 文件路径
            chunk_count: 文本块数量
            
        Returns:
            元数据字典
        """
        try:
            stat = file_path.stat()
            
            return {
                'filename': file_path.name,
                'file_path': str(file_path),
                'file_extension': file_path.suffix.lower(),
                'file_size': stat.st_size,
                'file_size_formatted': format_file_size(stat.st_size),
                'file_hash': get_file_hash(file_path),
                'created_time': stat.st_ctime,
                'modified_time': stat.st_mtime,
                'chunk_count': chunk_count,
                'chunk_size': self.chunk_size,
                'chunk_overlap': self.chunk_overlap
            }
        
        except Exception as e:
            logger.error(f"元数据生成失败: {file_path}, 错误: {e}")
            return {
                'filename': file_path.name,
                'file_path': str(file_path),
                'error': str(e)
            }
    
    def get_supported_formats(self) -> List[str]:
        """获取支持的文件格式列表"""
        return self.supported_formats.copy()
    
    def update_chunk_config(self, chunk_size: Optional[int] = None, chunk_overlap: Optional[int] = None):
        """
        更新文本分块配置
        
        Args:
            chunk_size: 新的块大小
            chunk_overlap: 新的块重叠大小
        """
        if chunk_size is not None:
            self.chunk_size = chunk_size
            logger.info(f"文本块大小已更新为: {chunk_size}")
        
        if chunk_overlap is not None:
            self.chunk_overlap = chunk_overlap
            logger.info(f"文本块重叠已更新为: {chunk_overlap}")
    
    def __len__(self):
        """
        防止意外调用len()时出错，同时提供调试信息
        """
        import traceback
        logger.warning("警告：对DocumentProcessor对象调用了len()函数，这可能是一个错误。")
        logger.warning("调用堆栈：")
        for line in traceback.format_stack():
            logger.warning(line.strip())
        return 0


# 全局文档处理器实例
document_processor = DocumentProcessor()